import json
import os
import random
import re
from dataclasses import dataclass
from io import BytesIO, StringIO
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

APP_NAME = "Model-Waldo"
DEFAULT_MODEL = "gpt-5.4-mini"
FALLBACK_MODEL = "gpt-5-mini"
MAX_SAMPLE_WORDS = 5000
FAST_SAMPLE_WORDS = 2000
LARGE_WORD_THRESHOLD = 8000
LARGE_CHAR_THRESHOLD = 40000
LARGE_ROW_THRESHOLD = 500
LARGE_PAGE_THRESHOLD = 15

CONTENT_TYPES = [
    "ui_strings",
    "marketing",
    "legal",
    "financial",
    "medical",
    "support",
    "ecommerce_catalog",
    "technical_documentation",
    "training_content",
    "general_business",
    "creative_literary",
    "mixed",
]

CONTENT_TYPE_OPTIONS = [
    "Auto-detect from source content",
    "UI strings",
    "Marketing",
    "Legal / financial",
    "Customer support",
    "Technical documentation",
    "Regulated industry",
    "General business",
]

CONTENT_TYPE_TO_INTERNAL = {
    "Auto-detect from source content": "auto_detect",
    "UI strings": "ui_strings",
    "Marketing": "marketing",
    "Legal / financial": "legal_financial",
    "Customer support": "support",
    "Technical documentation": "technical_documentation",
    "Regulated industry": "regulated_industry",
    "General business": "general_business",
}

TARGET_LANGUAGE_OPTIONS = [
    "Italian", "German", "French", "Spanish", "Portuguese", "Dutch",
    "Polish", "Turkish", "Arabic", "Hebrew", "Japanese", "Korean", "Chinese",
    "Hindi", "Vietnamese", "Thai", "Russian", "Ukrainian", "Swedish", "Danish",
    "Norwegian", "Finnish", "Other"
]

SOURCE_LANGUAGE = "English"

def render_xai_inspect_panel(analysis: dict, metadata: dict, ranked_df: pd.DataFrame, top_provider: pd.Series):
    """
    Renders an expandable deep-dive diagnostic dashboard breaking down the 
    underlying heuristic logic and scoring criteria behind Model-Waldo's output.
    """
    st.markdown("---")
    with st.expander("🕵️‍♂️ Model-Waldo System Inspection Panel (XAI)", expanded=True):
        st.markdown("### 🔬 Strategic Algorithmic Trace Overview")
        st.caption("This administrative telemetry dashboard breaks down how the current weights and prompt parameters yielded the active recommendation matrix.")

        # Tab Layout for Clean Telemetry Organization
        tab_token, tab_weights, tab_matrix = st.tabs([
            "🧠 Structured LLM Schema", 
            "⚖️ Core Metadata Metrics", 
            "📊 Provider Score Matrix"
        ])

        with tab_token:
            st.markdown("#### Raw Extracted Model Output Variables")
            st.json(analysis)
            st.caption("This data structure represents the direct type-safetied contract returned via the custom GPT schema wrapper.")

        with tab_weights:
            st.markdown("#### Syntactic Document Context Profiles")
            col1, col2 = st.columns(2)
            with col1:
                st.metric(label="Total Analyzed Character Footprint", value=f"{metadata.get('char_count', 0)} chars")
                st.metric(label="Detected Structural UI Placeholders", value=metadata.get('placeholder_count', 0))
            with col2:
                st.metric(label="File Extraction Source Paradigm", value=str(metadata.get('file_type', 'Raw Copy-Paste Input')))
                st.metric(label="Target Extraction Word Count Sample Limit", value=metadata.get('llm_sample_word_limit', 'Not Scoped'))

        with tab_matrix:
            st.markdown("#### Comprehensive Provider Metric Evaluations")
            st.dataframe(
                ranked_df[['Provider', 'Score', 'Base Quality Score', 'Dynamic Multiplier applied', 'Match Strategy']].sort_values(by="Score", ascending=False),
                use_container_width=True,
                hide_index=True
            )
            st.info(
                f"**Engine Routing Trace Summary:** `{top_provider['Provider']}` out-indexed competing localization nodes due to optimization matching for target parameters."
            )
ANALYSIS_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "primary_content_type": {"type": "string", "enum": CONTENT_TYPES},
        "secondary_content_types": {
            "type": "array",
            "items": {"type": "string", "enum": CONTENT_TYPES},
            "maxItems": 3,
        },
        "content_type_confidence": {"type": "number", "minimum": 0, "maximum": 1},
        "content_type_evidence": {"type": "array", "items": {"type": "string"}, "maxItems": 5},
        "source_content_complexity": {"type": "string", "enum": ["low", "medium", "high"]},
        "sentence_complexity": {"type": "string", "enum": ["low", "medium", "high"]},
        "idiom_slang_density": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "creative_language_density": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "cultural_nuance_level": {"type": "string", "enum": ["low", "medium", "high"]},
        "discipline_specific_jargon": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "ambiguity_level": {"type": "string", "enum": ["low", "medium", "high"]},
        "requires_transcreation": {"type": "boolean"},
        "domain_risk_level": {"type": "string", "enum": ["low", "medium", "high", "critical"]},
        "regulated_domain_detected": {"type": "string", "enum": ["none", "legal", "medical", "financial", "privacy", "safety", "insurance", "other"]},
        "brand_sensitivity": {"type": "string", "enum": ["low", "medium", "high"]},
        "tone_sensitivity": {"type": "string", "enum": ["low", "medium", "high"]},
        "tone_style": {"type": "string", "enum": ["neutral", "formal", "friendly", "technical", "persuasive", "luxury", "playful", "urgent", "empathetic", "mixed"]},
        "terminology_sensitivity": {"type": "string", "enum": ["low", "medium", "high"]},
        "product_terms_detected": {"type": "boolean"},
        "technical_terms_detected": {"type": "boolean"},
        "glossary_recommended": {"type": "boolean"},
        "term_consistency_risk": {"type": "string", "enum": ["low", "medium", "high"]},
        "placeholder_risk": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "html_or_xml_tag_risk": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "ui_truncation_risk": {"type": "string", "enum": ["none", "low", "medium", "high"]},
        "format_preservation_required": {"type": "boolean"},
        "locale_sensitivity": {"type": "string", "enum": ["low", "medium", "high"]},
        "cultural_adaptation_required": {"type": "boolean"},
        "contains_currency_dates_units": {"type": "boolean"},
        "contains_cultural_references": {"type": "boolean"},
        "contains_humor_or_taboo_risk": {"type": "boolean"},
        "privacy_sensitivity": {"type": "string", "enum": ["low", "medium", "high", "critical"]},
        "pii_detected": {"type": "boolean"},
        "confidential_business_content": {"type": "boolean"},
        "safe_for_public_model_api": {"type": "string", "enum": ["true", "false", "uncertain"]},
        "human_review_required": {"type": "boolean"},
        "human_review_type": {"type": "string", "enum": ["none", "linguistic_review", "subject_matter_review", "legal_review", "medical_review", "financial_compliance_review", "brand_review", "engineering_review"]},
        "human_review_reason": {"type": "string"},
        "key_detected_issues": {"type": "array", "items": {"type": "string"}, "maxItems": 8},
        "routing_recommendation_notes": {"type": "string"},
    },
    "required": [
        "primary_content_type", "secondary_content_types", "content_type_confidence",
        "content_type_evidence", "source_content_complexity", "sentence_complexity",
        "idiom_slang_density", "creative_language_density", "cultural_nuance_level",
        "discipline_specific_jargon", "ambiguity_level", "requires_transcreation",
        "domain_risk_level", "regulated_domain_detected", "brand_sensitivity", "tone_sensitivity",
        "tone_style", "terminology_sensitivity", "product_terms_detected", "technical_terms_detected",
        "glossary_recommended", "term_consistency_risk", "placeholder_risk", "html_or_xml_tag_risk",
        "ui_truncation_risk", "format_preservation_required", "locale_sensitivity",
        "cultural_adaptation_required", "contains_currency_dates_units", "contains_cultural_references",
        "contains_humor_or_taboo_risk", "privacy_sensitivity", "pii_detected", "confidential_business_content",
        "safe_for_public_model_api", "human_review_required", "human_review_type", "human_review_reason",
        "key_detected_issues", "routing_recommendation_notes"
    ],
}


def get_api_key() -> Optional[str]:
    # st.secrets raises an exception if no secrets.toml exists locally.
    try:
        if "OPENAI_API_KEY" in st.secrets:
            return st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return os.getenv("OPENAI_API_KEY")


def read_txt(uploaded_file) -> str:
    raw = uploaded_file.read()
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def read_docx(uploaded_file) -> str:
    document = Document(uploaded_file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(parts)


def read_pdf(uploaded_file) -> Tuple[str, int]:
    reader = PdfReader(uploaded_file)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages), len(reader.pages)


def extract_strings_from_json(obj: Any, path: str = "") -> List[str]:
    strings = []
    if isinstance(obj, dict):
        for key, value in obj.items():
            new_path = f"{path}.{key}" if path else str(key)
            strings.extend(extract_strings_from_json(value, new_path))
    elif isinstance(obj, list):
        for i, value in enumerate(obj):
            strings.extend(extract_strings_from_json(value, f"{path}[{i}]"))
    elif isinstance(obj, str):
        strings.append(f"{path}: {obj}" if path else obj)
    return strings


def read_json(uploaded_file) -> str:
    text = read_txt(uploaded_file)
    try:
        data = json.loads(text)
        return "\n".join(extract_strings_from_json(data))
    except json.JSONDecodeError:
        return text


def read_html(uploaded_file) -> str:
    html = read_txt(uploaded_file)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def dataframe_to_text(df: pd.DataFrame, source_column: str, max_rows: int = 5000) -> str:
    clean = df[source_column].dropna().astype(str).head(max_rows)
    return "\n".join(clean.tolist())


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))


def detect_placeholders(text: str) -> List[str]:
    patterns = [
        r"\{\{[^}]+\}\}", r"\{[^}\n]+\}", r"%[sdif]", r"\$\{[^}]+\}",
        r"%\([^)]+\)[sdif]", r"<%=?[^%]+%>", r"\[[A-Z0-9_\-]+\]"
    ]
    found = []
    for pattern in patterns:
        found.extend(re.findall(pattern, text))
    return sorted(set(found))[:50]


def detect_tags(text: str) -> List[str]:
    tags = re.findall(r"<\/?[A-Za-z][^>]*>", text)
    return sorted(set(tags))[:50]


def candidate_repeated_terms(text: str) -> List[str]:
    # Supporting signal only. Semantic interpretation is done by the LLM.
    words = re.findall(r"\b[A-Za-z][A-Za-z0-9\-]{3,}\b", text)
    stop = {
        "this", "that", "with", "from", "your", "have", "will", "when", "then", "there", "their", "about",
        "which", "into", "more", "than", "they", "them", "were", "been", "also", "only", "data", "user",
        "click", "open", "save", "cancel", "continue", "please", "error"
    }
    counts: Dict[str, int] = {}
    for word in words:
        key = word.strip()
        if key.lower() in stop:
            continue
        counts[key] = counts.get(key, 0) + 1
    terms = sorted([(k, v) for k, v in counts.items() if v >= 2], key=lambda x: (-x[1], x[0]))
    return [f"{k} ({v})" for k, v in terms[:20]]


def make_sample(text: str, max_words: int = MAX_SAMPLE_WORDS) -> str:
    words = text.split()
    if len(words) <= max_words:
        return text
    chunk = max_words // 3
    start = words[:chunk]
    middle_start = max(0, (len(words) // 2) - (chunk // 2))
    middle = words[middle_start:middle_start + chunk]
    end = words[-chunk:]
    return " ".join(start) + "\n\n[...MIDDLE SAMPLE...]\n\n" + " ".join(middle) + "\n\n[...END SAMPLE...]\n\n" + " ".join(end)


def language_pair_complexity(source: str, target: str) -> str:
    high = {"Japanese", "Korean", "Chinese", "Arabic", "Hebrew", "Thai", "Hindi", "Turkish", "Finnish", "Hungarian"}
    medium = {"German", "Russian", "Ukrainian", "Polish", "Dutch", "Swedish", "Danish", "Norwegian", "Vietnamese"}
    if source == target:
        return "low"
    if source in high or target in high:
        return "high"
    if source in medium or target in medium:
        return "medium"
    return "low"


def build_metadata(text: str, file_name: str = "pasted_text", file_type: str = "text", rows: Optional[int] = None, pages: Optional[int] = None) -> Dict[str, Any]:
    placeholders = detect_placeholders(text)
    tags = detect_tags(text)
    wc = word_count(text)
    char_count = len(text)
    large = wc > LARGE_WORD_THRESHOLD or char_count > LARGE_CHAR_THRESHOLD
    if rows is not None:
        large = large or rows > LARGE_ROW_THRESHOLD
    if pages is not None:
        large = large or pages > LARGE_PAGE_THRESHOLD
    return {
        "file_name": file_name,
        "file_type": file_type,
        "word_count": wc,
        "character_count": char_count,
        "row_count": rows,
        "page_count": pages,
        "is_large_file": large,
        "placeholder_count": len(placeholders),
        "placeholders_sample": placeholders[:10],
        "html_xml_tag_count": len(tags),
        "tags_sample": tags[:10],
        "repeated_candidate_terms": candidate_repeated_terms(text),
    }


def fallback_analysis(text: str, metadata: Dict[str, Any], known_domain: str) -> Dict[str, Any]:
    lower = text.lower()
    legal_terms = ["contract", "liability", "warranty", "terms", "compliance", "regulation", "claim", "policy"]
    medical_terms = ["dosage", "patient", "symptom", "treatment", "diagnosis", "clinical", "medical"]
    financial_terms = ["loan", "interest", "investment", "payment", "account", "tax", "fee", "credit"]
    marketing_terms = ["discover", "limited time", "exclusive", "transform", "unlock", "premium", "brand", "campaign"]
    ui_terms = ["save", "cancel", "submit", "settings", "login", "logout", "error", "password"]

    def any_terms(terms: List[str]) -> bool:
        return any(term in lower for term in terms)

    primary = "general_business"
    forced_regulated_context = False
    if known_domain and known_domain not in {"auto_detect", "Auto-detect from source content"}:
        if known_domain == "legal_financial":
            primary = "legal" if any_terms(legal_terms) or not any_terms(financial_terms) else "financial"
        elif known_domain == "regulated_industry":
            forced_regulated_context = True
            primary = "general_business"
        elif known_domain in CONTENT_TYPES:
            primary = known_domain
    elif metadata["placeholder_count"] > 0 or sum(1 for t in ui_terms if t in lower) >= 2:
        primary = "ui_strings"
    elif any_terms(legal_terms):
        primary = "legal"
    elif any_terms(medical_terms):
        primary = "medical"
    elif any_terms(financial_terms):
        primary = "financial"
    elif any_terms(marketing_terms):
        primary = "marketing"

    regulated = "none"
    risk = "low"
    review_type = "none"
    if forced_regulated_context:
        regulated = "other"
        risk = "high"
        review_type = "subject_matter_review"
    elif primary in {"legal", "medical", "financial"}:
        regulated = primary
        risk = "high"
        review_type = "legal_review" if primary == "legal" else "medical_review" if primary == "medical" else "financial_compliance_review"
    elif primary == "marketing":
        risk = "medium"
        review_type = "brand_review"

    return {
        "primary_content_type": primary,
        "secondary_content_types": [],
        "content_type_confidence": 0.55,
        "content_type_evidence": ["Fallback rule-based analysis was used because LLM analysis was unavailable."],
        "source_content_complexity": "medium",
        "sentence_complexity": "medium",
        "idiom_slang_density": "low",
        "creative_language_density": "medium" if primary == "marketing" else "low",
        "cultural_nuance_level": "medium" if primary == "marketing" else "low",
        "discipline_specific_jargon": "high" if primary in {"legal", "medical", "financial", "technical_documentation"} else "low",
        "ambiguity_level": "medium",
        "requires_transcreation": primary == "marketing",
        "domain_risk_level": risk,
        "regulated_domain_detected": regulated,
        "brand_sensitivity": "high" if primary == "marketing" else "low",
        "tone_sensitivity": "high" if primary == "marketing" else "medium",
        "tone_style": "persuasive" if primary == "marketing" else "neutral",
        "terminology_sensitivity": "high" if primary in {"ui_strings", "legal", "medical", "financial", "technical_documentation"} else "medium",
        "product_terms_detected": primary in {"ui_strings", "ecommerce_catalog", "technical_documentation"},
        "technical_terms_detected": primary in {"technical_documentation", "ui_strings"},
        "glossary_recommended": primary in {"ui_strings", "legal", "medical", "financial", "technical_documentation"},
        "term_consistency_risk": "high" if primary in {"ui_strings", "legal", "medical", "financial", "technical_documentation"} else "medium",
        "placeholder_risk": "high" if metadata["placeholder_count"] > 5 else "medium" if metadata["placeholder_count"] > 0 else "none",
        "html_or_xml_tag_risk": "medium" if metadata["html_xml_tag_count"] > 0 else "none",
        "ui_truncation_risk": "high" if primary == "ui_strings" else "none",
        "format_preservation_required": metadata["placeholder_count"] > 0 or metadata["html_xml_tag_count"] > 0,
        "locale_sensitivity": "high" if primary == "marketing" else "medium",
        "cultural_adaptation_required": primary == "marketing",
        "contains_currency_dates_units": bool(re.search(r"[$€£¥]|\b\d{1,2}/\d{1,2}/\d{2,4}\b|\bkg\b|\blbs\b", text)),
        "contains_cultural_references": primary == "marketing",
        "contains_humor_or_taboo_risk": False,
        "privacy_sensitivity": "medium" if re.search(r"[\w.%-]+@[\w.-]+\.[A-Za-z]{2,}|\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b", text) else "low",
        "pii_detected": bool(re.search(r"[\w.%-]+@[\w.-]+\.[A-Za-z]{2,}|\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b", text)),
        "confidential_business_content": False,
        "safe_for_public_model_api": "uncertain",
        "human_review_required": review_type != "none" or primary == "ui_strings",
        "human_review_type": review_type if review_type != "none" else "engineering_review" if primary == "ui_strings" else "none",
        "human_review_reason": "Fallback analysis detected content that should not go live without review." if review_type != "none" or primary == "ui_strings" else "No high-risk review trigger detected by fallback analysis.",
        "key_detected_issues": ["Fallback analysis only", f"Detected content type: {primary}", f"Placeholder count: {metadata['placeholder_count']}"],
        "routing_recommendation_notes": "Use fallback recommendation with caution. Add an OpenAI API key for stronger content analysis.",
    }


def call_openai_analysis(
    text_sample: str,
    metadata: Dict[str, Any],
    source_lang: str,
    target_lang: str,
    known_domain: str,
    content_requirements: Dict[str, bool],
    model_name: str,
    api_key: str,
) -> Dict[str, Any]:
    client = OpenAI(api_key=api_key)
    system_prompt = """
You are a senior localization solutions architect. Analyze source localization content for provider routing.
Return only structured JSON that matches the provided schema. Do not translate the content.
Do not recommend a provider. Your job is content analysis only. The local scoring engine makes the provider decision.
Be conservative about human review for regulated, safety, medical, legal, financial, brand-sensitive, or complex terminology content.
Use evidence from the content and metadata. If the user provided a content type or content requirements, treat them as important operational signals.
""".strip()

    user_payload = {
        "source_language": source_lang,
        "target_language": target_lang,
        "selected_content_type_or_auto_detect": known_domain,
        "content_requirements": content_requirements,
        "metadata": metadata,
        "source_content_sample": text_sample,
        "instruction": "Analyze the source content for localization routing. Return structured JSON only. Do not translate.",
    }

    try:
        response = client.responses.create(
            model=model_name,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
            ],
            text={
                "format": {
                    "type": "json_schema",
                    "name": "model_waldo_content_analysis",
                    "schema": ANALYSIS_SCHEMA,
                    "strict": True,
                }
            },
        )
        return json.loads(response.output_text)
    except Exception as first_error:
        # Retry once with fallback model if the selected model fails.
        if model_name != FALLBACK_MODEL:
            try:
                response = client.responses.create(
                    model=FALLBACK_MODEL,
                    input=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
                    ],
                    text={
                        "format": {
                            "type": "json_schema",
                            "name": "model_waldo_content_analysis",
                            "schema": ANALYSIS_SCHEMA,
                            "strict": True,
                        }
                    },
                )
                return json.loads(response.output_text)
            except Exception as second_error:
                raise RuntimeError(f"OpenAI analysis failed with {model_name}: {first_error}; fallback {FALLBACK_MODEL} also failed: {second_error}")
        raise RuntimeError(f"OpenAI analysis failed: {first_error}")


def load_provider_data() -> pd.DataFrame:
    return pd.read_csv("data/providers.csv")


def level_to_score(value: str, low: int = 1, medium: int = 3, high: int = 5, critical: int = 5) -> int:
    mapping = {"none": 0, "low": low, "medium": medium, "high": high, "critical": critical, "true": 5, "false": 1, "uncertain": 3}
    return mapping.get(str(value).lower(), medium)


def target_expansion_factor(target: str) -> float:
    """Rough output-length factor for estimating LLM output tokens. Demo estimate only."""
    high_expansion = {"German", "Finnish", "Turkish", "Russian", "Ukrainian", "Polish"}
    moderate_expansion = {"Italian", "French", "Spanish", "Portuguese", "Dutch", "Swedish", "Danish", "Norwegian"}
    compact = {"Chinese", "Japanese", "Korean"}
    if target in high_expansion:
        return 1.18
    if target in moderate_expansion:
        return 1.10
    if target in compact:
        return 0.85
    return 1.05


def estimate_provider_cost(provider_row: pd.Series, metadata: Dict[str, Any], target_lang: str) -> Tuple[float, float, str]:
    """Return estimated job cost, cost per 1k source words, and pricing basis.

    This is a normalized MVP estimate, not an invoice estimate. MT providers are priced
    by source characters. LLM providers are estimated from input and output tokens.
    """
    chars = max(1, int(metadata.get("character_count", 0) or 0))
    words = max(1, int(metadata.get("word_count", 0) or 0))
    pricing_basis = str(provider_row.get("pricing_basis", "profile"))

    if pricing_basis == "chars":
        rate = float(provider_row.get("usd_per_million_chars", 0) or 0)
        cost = (chars / 1_000_000) * rate
        basis_label = f"chars @ ${rate:g}/1M chars"
    elif pricing_basis == "tokens":
        input_rate = float(provider_row.get("usd_per_million_input_tokens", 0) or 0)
        output_rate = float(provider_row.get("usd_per_million_output_tokens", 0) or 0)
        input_tokens = chars / 4.0
        output_tokens = input_tokens * target_expansion_factor(target_lang)
        cost = (input_tokens / 1_000_000) * input_rate + (output_tokens / 1_000_000) * output_rate
        basis_label = f"tokens @ ${input_rate:g}/1M in + ${output_rate:g}/1M out"
    else:
        # Fallback to old coarse profile if pricing columns are missing.
        profile = str(provider_row.get("cost_profile", "medium"))
        pseudo_rate = {"low": 8, "medium": 18, "high": 60}.get(profile, 18)
        cost = (chars / 1_000_000) * pseudo_rate
        basis_label = f"demo profile: {profile}"

    return round(cost, 6), round((cost / words) * 1000, 6), basis_label


def normalized_cost_score(cost: float, min_cost: float, max_cost: float) -> float:
    """Higher score is better. Cheapest gets 5.0, highest gets 1.0."""
    if max_cost <= min_cost:
        return 3.5
    score = 5 - ((cost - min_cost) / (max_cost - min_cost)) * 4
    return max(1.0, min(5.0, score))


def latency_score(latency_profile: str, urgency: str) -> float:
    # Low latency = fast = better score. Urgency is the only user input that boosts speed.
    base = {"low": 5, "medium": 3.5, "high": 2}.get(latency_profile, 3)
    if urgency == "High":
        return base
    if urgency == "Low":
        return min(5, base + 0.25)
    return base


def content_fit_column(primary_content_type: str) -> str:
    mapping = {
        "ui_strings": "ui_fit",
        "marketing": "marketing_fit",
        "legal": "legal_fit",
        "financial": "financial_fit",
        "medical": "medical_fit",
        "support": "support_fit",
        "ecommerce_catalog": "ecommerce_fit",
        "technical_documentation": "technical_fit",
        "training_content": "technical_fit",
        "creative_literary": "marketing_fit",
        "general_business": "support_fit",
        "mixed": "quality_general",
    }
    return mapping.get(primary_content_type, "quality_general")


def dynamic_weights(priority: str, content_requirements: Optional[Dict[str, bool]] = None) -> Dict[str, float]:
    weights = {
        "quality_fit": 0.35,
        "language_pair_fit": 0.20,
        "cost_fit": 0.15,
        "latency_fit": 0.10,
        "terminology_fit": 0.15,
        "privacy_fit": 0.05,
    }
    boosts = {
        "Quality": "quality_fit",
        "Cost": "cost_fit",
    }
    if priority in boosts:
        weights[boosts[priority]] += 0.15

    reqs = content_requirements or {}
    if reqs.get("glossary_terminology_file_required"):
        weights["terminology_fit"] += 0.10
    if reqs.get("brand_tone_transcreation_required"):
        weights["quality_fit"] += 0.10
    if reqs.get("placeholder_tag_format_preservation_required"):
        weights["terminology_fit"] += 0.05

    total = sum(weights.values())
    return {k: v / total for k, v in weights.items()}


def language_fit(
    provider_row: pd.Series,
    pair_complexity: str,
    source: str,
    target: str,
    analysis: Optional[Dict[str, Any]] = None,
    content_requirements: Optional[Dict[str, bool]] = None,
) -> float:
    """Contextual language fit, not a static language-pair score.

    Baseline language strength still matters, but content type and requirements modify
    the score so one strong EN-DE provider does not dominate every scenario.
    """
    european = {"English", "Italian", "German", "French", "Spanish", "Portuguese", "Dutch", "Polish", "Swedish", "Danish", "Norwegian"}
    if source in european and target in european:
        base = float(provider_row["language_strength_european"])
    elif pair_complexity == "high":
        base = float(provider_row["language_strength_high_distance"])
    else:
        base = float(provider_row["language_strength_general"])

    analysis = analysis or {}
    reqs = content_requirements or {}
    primary = str(analysis.get("primary_content_type", "general_business"))
    is_llm = str(provider_row.get("workflow_type", "")).lower() == "llm"
    is_mt = str(provider_row.get("workflow_type", "")).lower() == "mt"
    transcreation_required = bool(analysis.get("requires_transcreation", False)) or reqs.get("brand_tone_transcreation_required", False)

    if primary in {"marketing", "creative_literary"} or transcreation_required:
        if is_llm:
            base += 0.55
        elif is_mt:
            base -= 0.30

    if primary == "ui_strings":
        if is_mt:
            base += 0.15
        elif is_llm and str(provider_row.get("latency_profile", "")).lower() == "high":
            base -= 0.15

    if primary in {"legal", "financial", "medical"}:
        # Keep language fit important, but let domain fit and review policy carry the safety decision.
        base += min(0.15, (float(provider_row.get(content_fit_column(primary), 3)) - 3.5) / 10)

    return max(1.0, min(5.0, base))


def score_providers(
    providers: pd.DataFrame,
    analysis: Dict[str, Any],
    source_lang: str,
    target_lang: str,
    priority: str,
    urgency: str,
    pair_complexity: str,
    content_requirements: Optional[Dict[str, bool]] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> pd.DataFrame:
    weights = dynamic_weights(priority, content_requirements)
    fit_col = content_fit_column(analysis["primary_content_type"])
    terminology_need = level_to_score(analysis["terminology_sensitivity"])
    privacy_need = level_to_score(analysis["privacy_sensitivity"])
    metadata = metadata or {"character_count": 0, "word_count": 1}

    cost_estimates = []
    for _, provider in providers.iterrows():
        est_cost, est_cost_per_1k_words, pricing_basis = estimate_provider_cost(provider, metadata, target_lang)
        cost_estimates.append((est_cost, est_cost_per_1k_words, pricing_basis))
    costs_only = [c[0] for c in cost_estimates]
    min_cost = min(costs_only) if costs_only else 0
    max_cost = max(costs_only) if costs_only else 0

    rows = []
    for i, (_, provider) in enumerate(providers.iterrows()):
        quality = (float(provider["quality_general"]) + float(provider[fit_col])) / 2
        lang = language_fit(provider, pair_complexity, source_lang, target_lang, analysis, content_requirements)
        est_cost, est_cost_per_1k_words, pricing_basis = cost_estimates[i]
        cost = normalized_cost_score(est_cost, min_cost, max_cost)
        latency = latency_score(provider["latency_profile"], urgency)

        # Glossary support matters most when terminology sensitivity is high.
        glossary = float(provider["glossary_support"])
        terminology = (glossary * (terminology_need / 5)) + (float(provider["format_safety"]) * (1 - terminology_need / 5))

        # Privacy controls remain a minor profile field, but privacy is not a user business priority in the MVP.
        privacy = float(provider["privacy_profile"])
        privacy_fit = privacy if privacy_need >= 3 else min(5, privacy + 0.5)

        primary = analysis["primary_content_type"]
        is_llm = str(provider.get("workflow_type", "")).lower() == "llm"
        is_mt = str(provider.get("workflow_type", "")).lower() == "mt"
        brand_transcreation_required = bool(analysis.get("requires_transcreation", False)) or bool((content_requirements or {}).get("brand_tone_transcreation_required"))

        # Brand/transcreation content should not be dominated by static EN-DE language fit.
        contextual_bonus = 0.0
        if primary in {"marketing", "creative_literary"} or brand_transcreation_required:
            if is_llm:
                contextual_bonus += 0.35
                quality = max(quality, float(provider["transcreation_fit"]) - 0.05)
            elif is_mt:
                contextual_bonus -= 0.15

        if primary == "ui_strings" and urgency == "High" and priority == "Cost":
            if is_mt:
                contextual_bonus += 0.10
            elif is_llm and provider["cost_profile"] == "low" and provider["latency_profile"] == "low":
                contextual_bonus += 0.05

        raw_score = (
            quality * weights["quality_fit"]
            + lang * weights["language_pair_fit"]
            + cost * weights["cost_fit"]
            + latency * weights["latency_fit"]
            + terminology * weights["terminology_fit"]
            + privacy_fit * weights["privacy_fit"]
            + contextual_bonus
        )

        penalty = 0.0
        if analysis["domain_risk_level"] in {"high", "critical"}:
            # Penalize light/low-quality automated options more than strong engines. Human review is handled separately.
            if float(provider[fit_col]) < 3.5:
                penalty += 0.30
            if provider["cost_profile"] == "low" and priority != "Cost":
                penalty += 0.15
        if brand_transcreation_required and float(provider["transcreation_fit"]) < 4:
            penalty += 0.55
        if analysis["placeholder_risk"] in {"medium", "high"} and float(provider["format_safety"]) < 4:
            penalty += 0.25

        final_score = max(0, min(100, (raw_score - penalty) * 20))

        rows.append({
            "Provider": provider["provider"],
            "Score": round(final_score, 1),
            "Quality fit": round(quality, 2),
            "Language fit": round(lang, 2),
            "Cost fit": round(cost, 2),
            "Estimated job cost": f"${est_cost:.4f}",
            "Est. cost / 1k words": f"${est_cost_per_1k_words:.4f}",
            "Pricing basis": pricing_basis,
            "Turnaround speed fit": round(latency, 2),
            "Terminology fit": round(terminology, 2),
            "Privacy fit": round(privacy_fit, 2),
            "Workflow type": provider["workflow_type"],
            "Evidence confidence": provider["evidence_confidence"],
            "Evidence notes": provider["evidence_notes"],
            "Evidence sources": provider["evidence_sources"],
        })
    return pd.DataFrame(rows).sort_values("Score", ascending=False).reset_index(drop=True)


def apply_content_type_context(analysis: Dict[str, Any], selected_content_type: str) -> Dict[str, Any]:
    """Apply the user's selected content type as an operational override when not auto-detect."""
    a = dict(analysis)
    if selected_content_type in {"auto_detect", "Auto-detect from source content", None, ""}:
        return a

    if selected_content_type in CONTENT_TYPES:
        if a.get("primary_content_type") != selected_content_type:
            a.setdefault("content_type_evidence", [])
            a["content_type_evidence"] = [f"User selected Content Type: {selected_content_type.replace('_', ' ')}."] + list(a.get("content_type_evidence", []))[:4]
        a["primary_content_type"] = selected_content_type
    elif selected_content_type == "legal_financial":
        if a.get("primary_content_type") not in {"legal", "financial"}:
            a["primary_content_type"] = "legal"
        a["domain_risk_level"] = "high"
        if a.get("regulated_domain_detected") == "none":
            a["regulated_domain_detected"] = "legal"
        a["glossary_recommended"] = True
        a["terminology_sensitivity"] = "high"
    elif selected_content_type == "regulated_industry":
        a["domain_risk_level"] = "high"
        if a.get("regulated_domain_detected") == "none":
            a["regulated_domain_detected"] = "other"
        a["human_review_required"] = True
        if a.get("human_review_type") == "none":
            a["human_review_type"] = "subject_matter_review"
    return a


def apply_content_requirements(analysis: Dict[str, Any], content_requirements: Dict[str, bool]) -> Dict[str, Any]:
    """Make user-declared operational requirements affect analysis and scoring."""
    a = dict(analysis)
    issues = list(a.get("key_detected_issues", []))

    if content_requirements.get("glossary_terminology_file_required"):
        a["glossary_recommended"] = True
        a["terminology_sensitivity"] = "high"
        a["term_consistency_risk"] = "high"
        issues.append("User indicated that a glossary / terminology file is required.")

    if content_requirements.get("placeholder_tag_format_preservation_required"):
        a["format_preservation_required"] = True
        if a.get("placeholder_risk") == "none":
            a["placeholder_risk"] = "medium"
        issues.append("User indicated that placeholder, tag, or format preservation is required.")

    if content_requirements.get("brand_tone_transcreation_required"):
        a["requires_transcreation"] = True
        a["brand_sensitivity"] = "high"
        a["tone_sensitivity"] = "high"
        issues.append("User indicated that brand tone / transcreation is required.")

    a["key_detected_issues"] = issues[:8]
    return a

def has_strong_legal_or_regulated_indicators(text: str) -> bool:
    """Avoid treating ordinary UI/security/account terms as legal content."""
    lower = text.lower()
    strong_patterns = [
        r"\bterms of service\b", r"\bterms and conditions\b", r"\bprivacy policy\b",
        r"\blegal disclaimer\b", r"\bliability\b", r"\bindemnif", r"\bwarranty\b",
        r"\bcontract\b", r"\bagreement\b", r"\bcompliance\b", r"\bregulat",
        r"\bgdpr\b", r"\bccpa\b", r"\bhipaa\b", r"\bfinancial risk\b",
        r"\binvestment risk\b", r"\bnot legal advice\b", r"\bnot financial advice\b",
        r"\bsubject to applicable law\b", r"\bgoverning law\b", r"\barbitration\b",
    ]
    return any(re.search(pattern, lower) for pattern in strong_patterns)


def apply_operational_review_policy(analysis: Dict[str, Any], metadata: Dict[str, Any], source_text: str = "") -> Dict[str, Any]:
    """Normalize LLM review recommendations into Model-Waldo's operating policy.

    The LLM analyzes content, but the app owns the final human-review guidance.
    Human review means review of the localized output, not review of the provider choice.
    """
    a = dict(analysis)
    primary = str(a.get("primary_content_type", "general_business"))
    risk = str(a.get("domain_risk_level", "low")).lower()
    regulated = str(a.get("regulated_domain_detected", "none")).lower()
    privacy = str(a.get("privacy_sensitivity", "low")).lower()
    brand = str(a.get("brand_sensitivity", "low")).lower()
    terminology = str(a.get("terminology_sensitivity", "low")).lower()
    placeholder_risk = str(a.get("placeholder_risk", "none")).lower()
    transcreation = bool(a.get("requires_transcreation", False))
    placeholder_count = int(metadata.get("placeholder_count", 0) or 0)
    tag_count = int(metadata.get("html_xml_tag_count", 0) or 0)
    strong_legal_indicators = has_strong_legal_or_regulated_indicators(source_text)

    required = False
    review_type = "none"
    reason = "No mandatory human review trigger detected. Normal localization QA still applies."
    qa_note = ""

    # UI strings are common sources of false positives. Treat ordinary billing/security/API labels as UI unless strong legal indicators are present.
    if primary == "ui_strings" and not strong_legal_indicators and risk != "critical" and regulated not in {"medical", "financial", "safety", "insurance"}:
        required = False
        review_type = "none"
        reason = "UI strings were detected, but no strong legal, regulated, or high-risk content was found. Mandatory human review is not required."
        qa_note = "Automated QA recommended: placeholder preservation, tag preservation, glossary consistency, and length/truncation checks."

    # Regulated and high-impact domains override review tolerance.
    elif primary in {"legal", "medical", "financial"} or regulated in {"legal", "medical", "financial", "safety", "insurance"} or (risk == "critical") or (risk == "high" and strong_legal_indicators):
        required = True
        if primary == "medical" or regulated == "medical":
            review_type = "medical_review"
            reason = "Medical or health-related content was detected. The localized output should be reviewed for safety, accuracy, and regulatory impact."
        elif primary == "financial" or regulated in {"financial", "insurance"}:
            review_type = "financial_compliance_review"
            reason = "Financial, insurance, or compliance-sensitive content was detected. The localized output should be reviewed for accuracy and regulatory risk."
        elif primary == "legal" or regulated == "legal" or strong_legal_indicators:
            review_type = "legal_review"
            reason = "Legal, contractual, privacy-policy, or compliance language was detected. The localized output should be reviewed before release because wording can affect legal meaning."
        else:
            review_type = "subject_matter_review"
            reason = "High-risk or regulated content was detected. The localized output should be reviewed by a qualified subject-matter reviewer."

    # Privacy sensitivity alone should not turn UI/account strings into legal review.
    elif privacy in {"high", "critical"} and strong_legal_indicators:
        required = True
        review_type = "legal_review"
        reason = "High privacy sensitivity and privacy-policy/compliance language were detected. Legal or compliance review is recommended."

    # Brand/transcreation content usually needs brand review unless the risk is clearly low.
    elif primary in {"marketing", "creative_literary"} or transcreation or brand == "high":
        required = True
        review_type = "brand_review"
        reason = "Brand-sensitive or creative content was detected. The localized output should be reviewed for tone, persuasion, cultural adaptation, and brand voice."

    elif privacy in {"high", "critical"}:
        required = True
        review_type = "subject_matter_review"
        reason = "High privacy sensitivity was detected. The localized output should be reviewed before release to reduce confidentiality and compliance risk."

    a["human_review_required"] = required
    a["human_review_type"] = review_type
    a["human_review_reason"] = reason
    if qa_note:
        a["operational_qa_note"] = qa_note
    elif "operational_qa_note" in a:
        a.pop("operational_qa_note", None)
    return a



def human_review_summary(analysis: Dict[str, Any]) -> str:
    if not analysis["human_review_required"]:
        msg = "Not mandatory. Normal QA still applies."
        if analysis.get("operational_qa_note"):
            msg += " " + analysis["operational_qa_note"]
        return msg
    return f"Required before release. {analysis['human_review_reason']}"


def title_case_signal(value: str) -> str:
    return str(value or "").replace("_", " ").strip().title()


def content_type_display(analysis: Dict[str, Any]) -> str:
    content_type = title_case_signal(analysis.get("primary_content_type", "unknown"))
    confidence = float(analysis.get("content_type_confidence", 0) or 0)
    if confidence < 0.85:
        return f"Based on the provided document, Model-Waldo believes this content is {content_type}."
    return content_type


def glossary_status(analysis: Dict[str, Any], content_requirements: Dict[str, bool]) -> str:
    if content_requirements.get("glossary_terminology_file_required"):
        return "Required"
    if analysis.get("glossary_recommended"):
        return "Recommended"
    return "Not required"


def format_status(analysis: Dict[str, Any], content_requirements: Dict[str, bool]) -> str:
    if content_requirements.get("placeholder_tag_format_preservation_required"):
        return "Required"
    if analysis.get("format_preservation_required") or analysis.get("placeholder_risk") in {"medium", "high"} or analysis.get("html_or_xml_tag_risk") in {"medium", "high"}:
        return "Recommended"
    return "No major concerns detected"


def strong_alternative(ranked: pd.DataFrame) -> str:
    if ranked.shape[0] < 2:
        return "None"
    return str(ranked.iloc[1]["Provider"])


def localized_output_review_label(analysis: Dict[str, Any]) -> str:
    if not analysis.get("human_review_required"):
        return "Not mandatory"
    return "Required"


def review_type_label(analysis: Dict[str, Any]) -> str:
    if not analysis.get("human_review_required"):
        return "Standard QA"
    return title_case_signal(analysis.get("human_review_type", "review"))


def qa_focus(analysis: Dict[str, Any], content_requirements: Dict[str, bool]) -> str:
    focus = []
    if analysis.get("terminology_sensitivity") in {"medium", "high"} or content_requirements.get("glossary_terminology_file_required"):
        focus.append("terminology consistency")
    if analysis.get("placeholder_risk") in {"medium", "high"} or analysis.get("html_or_xml_tag_risk") in {"medium", "high"} or content_requirements.get("placeholder_tag_format_preservation_required"):
        focus.append("placeholder / format preservation")
    if analysis.get("ui_truncation_risk") in {"medium", "high"}:
        focus.append("length / truncation checks")
    if analysis.get("human_review_type") == "legal_review":
        focus.append("legal accuracy")
    if analysis.get("human_review_type") == "brand_review" or content_requirements.get("brand_tone_transcreation_required"):
        focus.append("brand tone and cultural adaptation")
    if not focus:
        focus.append("standard localization QA")
    return ", ".join(dict.fromkeys(focus))


def short_review_reason(analysis: Dict[str, Any]) -> str:
    if not analysis.get("human_review_required"):
        if analysis.get("operational_qa_note"):
            return analysis["operational_qa_note"]
        return "No mandatory human review trigger detected. Normal QA still applies."
    rt = analysis.get("human_review_type")
    if rt == "legal_review":
        return "Legal meaning must be preserved. Wording can affect legal meaning."
    if rt == "financial_compliance_review":
        return "Financial or compliance meaning must be preserved before release."
    if rt == "medical_review":
        return "Safety and medical accuracy must be validated before release."
    if rt == "brand_review":
        return "Tone, persuasion, and cultural adaptation should be validated."
    if rt == "engineering_review":
        return "UI behavior, placeholders, layout, and truncation should be checked."
    return analysis.get("human_review_reason", "Review recommended before release.")


def build_why_bullets(top: pd.Series, ranked: pd.DataFrame, analysis: Dict[str, Any], pair_complexity: str, content_requirements: Dict[str, bool]) -> List[str]:
    bullets = [
        f"Recommended model is {top['Provider']} based on the combined content, language, urgency, and requirement fit.",
        f"Content type: {content_type_display(analysis)}.",
        f"English → target language has {pair_complexity} language-pair complexity.",
    ]
    if glossary_status(analysis, content_requirements) in {"Required", "Recommended"}:
        bullets.append(f"Glossary / terminology control is {glossary_status(analysis, content_requirements).lower()}.")
    if format_status(analysis, content_requirements) in {"Required", "Recommended"}:
        bullets.append(f"Format / placeholder preservation is {format_status(analysis, content_requirements).lower()}.")
    alt = strong_alternative(ranked)
    if alt != "None":
        bullets.append(f"Strong alternative: {alt}.")
    if analysis.get("human_review_required"):
        bullets.append(f"Localized output review is required: {review_type_label(analysis).lower()}.")
    elif analysis.get("operational_qa_note"):
        bullets.append(analysis["operational_qa_note"])
    return bullets


def render_decision_dashboard(
    top: pd.Series,
    ranked: pd.DataFrame,
    analysis: Dict[str, Any],
    pair_complexity: str,
    content_requirements: Dict[str, bool],
):
    st.header("2. Recommendation")

    # Extract confidence metric safely (default to low if completely missing)
    confidence = float(analysis.get("content_type_confidence", 0.50))
    confidence_percentage = int(confidence * 100)

    # Inject a dynamic visual buffer badge system based on model assurance
    if confidence >= 0.85:
        st.success(
            f"🎯 **High Confidence Routing Strategy Matrix:** Model-Waldo is **{confidence_percentage}%** confident in this classification profile. The risk coefficient is well within standard automated parameters.",
            icon="✅"
        )
    elif 0.70 <= confidence < 0.85:
        st.warning(
            f"⚠️ **Ambiguous Content Density Alert:** Model-Waldo is only **{confidence_percentage}%** confident in this structural parsing analysis. Review alternative model outputs below if the primary choice feels unaligned.",
            icon="⚡"
        )
    else:
        st.error(
            f"🚨 **Critical Classification Variance Warning:** Model-Waldo outputted a **{confidence_percentage}%** confidence floor for this sample dataset. Cross-referencing alternative providers in the collapsed matrix view below is strongly recommended.",
            icon="🛑"
        )

    # Create your core metric configuration layout
    rec_col, analysis_col, review_col = st.columns(3)

    with rec_col:
        # Determine border aesthetics based on confidence scoring thresholds
        border_status = "red" if confidence < 0.70 else ("orange" if confidence < 0.85 else "green")
        
        # Inject custom styling directly to draw eyes to structural instability
        with st.container(border=True):
            st.subheader("Recommended Model")
            st.markdown(f"## {top['Provider']}")
            
            if confidence < 0.70:
                st.caption(f"⚠️ *Review Alternative Providers below*")
            else:
                st.markdown("**Recommended**")
                
            st.divider()
            st.write(f"**Score:** {top['Score']} / 100")
            alt = strong_alternative(ranked)
            if alt != "None":
                st.write(f"**Strong alternative:** {alt}")

    with analysis_col:
        with st.container(border=True):
            st.subheader("Content Analysis")
            st.write(f"**Content type**  \n{content_type_display(analysis)}")
            st.write(f"**Language-pair complexity**  \n{title_case_signal(pair_complexity)}")
            st.write(f"**Glossary / terminology**  \n{glossary_status(analysis, content_requirements)}")
            st.write(f"**Format preservation**  \n{format_status(analysis, content_requirements)}")

    with review_col:
        with st.container(border=True):
            st.subheader("Review / QA")
            
            # If model confidence drops below baseline requirements, force manual review signals
            if confidence < 0.70:
                st.write(f"**Localized output review**  \n🔴 Escalated to Mandatory")
                st.write(f"**Review type**  \nLinguistic Peer Review")
                st.write(f"**Reason**  \nLow classification mapping confidence ({confidence_percentage}%). Verification required.")
            else:
                st.write(f"**Localized output review**  \n{localized_output_review_label(analysis)}")
                st.write(f"**Review type**  \n{review_type_label(analysis)}")
                st.write(f"**Reason**  \n{short_review_reason(analysis)}")
                
            st.write(f"**QA focus**  \n{qa_focus(analysis, content_requirements)}")

    # Summary overview layout stays transparent
    with st.container(border=True):
        st.subheader("Why this recommendation?")
        for item in build_why_bullets(top, ranked, analysis, pair_complexity, content_requirements):
            st.write(f"- {item}")


def render_content_analysis_details(analysis: Dict[str, Any], content_requirements: Dict[str, bool], target_lang: str):
    st.header("3. Content Analysis Detail")
    with st.container(border=True):
        st.write("**Content Type:**")
        st.write(f"- {content_type_display(analysis)}")

        st.write("**Evidence:**")
        evidence = analysis.get("content_type_evidence", []) or ["No classification evidence returned."]
        for item in evidence:
            st.write(f"- {item}")

        st.write("**Localization Risks:**")
        risks = list(analysis.get("key_detected_issues", []))
        if content_requirements.get("glossary_terminology_file_required") and "Glossary / terminology file is required." not in risks:
            risks.append("Glossary / terminology file is required.")
        if content_requirements.get("placeholder_tag_format_preservation_required") and "Placeholder, tag, or format preservation is required." not in risks:
            risks.append("Placeholder, tag, or format preservation is required.")
        if content_requirements.get("brand_tone_transcreation_required") and "Brand tone / transcreation is required." not in risks:
            risks.append("Brand tone / transcreation is required.")
        if target_lang == "German" and "German text expansion may affect constrained layouts." not in risks:
            risks.append("German text expansion may affect constrained layouts.")
        if analysis.get("human_review_required"):
            risks.append(f"Localized output review required: {review_type_label(analysis)}.")
        elif analysis.get("operational_qa_note"):
            risks.append(analysis["operational_qa_note"])

        for item in risks[:10]:
            st.write(f"- {item}")


def main():
    st.set_page_config(page_title=APP_NAME, page_icon="🔎", layout="wide")
    st.title("🔎 Model-Waldo")
    st.caption("Content-Aware MT/LLM Provider Selection Engine")

    with st.expander("What this prototype does", expanded=False):
        st.write(
            "Model-Waldo analyzes real English source content and business constraints, then ranks MT/LLM providers and model variants using explainable scoring rules. "
            "The LLM analyzes content. The local scoring engine makes the provider/model decision. The tool does not translate content."
        )
        st.write("Sample datasets are stored separately in the `/sample_data` folder. Upload them like any other file.")
        st.write("MVP supports English source content only. Multisource-language routing is planned for V2.")

    st.warning(
        "Privacy note: This MVP uses the OpenAI API for content analysis. Do not upload confidential, regulated, personal, or production-sensitive content unless you have the appropriate authorization and enterprise data controls in place.",
        icon="⚠️",
    )

    api_key = get_api_key()
    if not api_key:
        st.warning("No OpenAI API key detected. The app will run in fallback demo mode. Add OPENAI_API_KEY in Streamlit secrets or as an environment variable for LLM analysis.")

    providers = load_provider_data()

    with st.sidebar:
        st.header("Request Setup")
        st.text_input("Source language", value=SOURCE_LANGUAGE, disabled=True, help="MVP supports English source content only. Multisource routing is planned for V2.")
        source_lang = SOURCE_LANGUAGE
        target_lang = st.selectbox("Target language", TARGET_LANGUAGE_OPTIONS, index=TARGET_LANGUAGE_OPTIONS.index("German"))
        priority = st.selectbox("Business priority", ["None", "Quality", "Cost"], index=0, help="None applies no extra business-priority boost. Quality or Cost adjusts provider/model scoring.")
        urgency = st.selectbox("Urgency", ["Low", "Medium", "High"], index=1)
        selected_content_type = st.selectbox("Content Type", CONTENT_TYPE_OPTIONS, index=0)
        known_domain = CONTENT_TYPE_TO_INTERNAL[selected_content_type]
        st.markdown("**Content Requirements**")
        st.caption("Select only the requirements that are truly required for this project. The LLM will still detect additional risks automatically.")
        content_requirements = {
            "glossary_terminology_file_required": st.checkbox("Glossary / terminology file required", value=False),
            "placeholder_tag_format_preservation_required": st.checkbox("Placeholder, tag, or format preservation required", value=False),
            "brand_tone_transcreation_required": st.checkbox("Brand tone / transcreation required", value=False),
        }
        model_name = st.text_input("OpenAI model for content analysis", value=DEFAULT_MODEL)
        fast_mode = st.checkbox("Fast demo mode", value=False, help="Caps the LLM sample closer to 2,000 words.")
# --- EXPLAINABLE AI (XAI) SIDEBAR CONFIGURATION ---
st.sidebar.markdown("---")
st.sidebar.subheader("🛠️ Developer Sandbox")
dev_inspect_mode = st.sidebar.toggle(
    "Enable Inspect Mode (XAI)", 
    value=False,
    help="Surfaces internal calculation matrices, structural token allocations, and raw model schema parameters live."
)
    st.header("1. Add Source Content")
    upload_tab, paste_tab = st.tabs(["Upload file", "Paste text"])

    extracted_text = ""
    metadata: Dict[str, Any] = {}
    file_name = "pasted_text"
    file_type = "text"

    with upload_tab:
        uploaded_file = st.file_uploader(
            "Upload source content",
            type=["txt", "csv", "xlsx", "docx", "pdf", "json", "html", "htm"],
            help="MVP supports text-based files only. Scanned PDFs and OCR are out of scope.",
        )
        if uploaded_file:
            file_name = uploaded_file.name
            ext = file_name.split(".")[-1].lower()
            file_type = ext
            rows = None
            pages = None
            try:
                if ext == "txt":
                    extracted_text = read_txt(uploaded_file)
                elif ext == "csv":
                    df = pd.read_csv(uploaded_file)
                    rows = len(df)
                    st.write("Preview")
                    st.dataframe(df.head(10), use_container_width=True)
                    source_column = st.selectbox("Which column contains the source text?", df.columns.tolist())
                    extracted_text = dataframe_to_text(df, source_column)
                elif ext == "xlsx":
                    df = pd.read_excel(uploaded_file)
                    rows = len(df)
                    st.write("Preview")
                    st.dataframe(df.head(10), use_container_width=True)
                    source_column = st.selectbox("Which column contains the source text?", df.columns.tolist())
                    extracted_text = dataframe_to_text(df, source_column)
                elif ext == "docx":
                    extracted_text = read_docx(uploaded_file)
                elif ext == "pdf":
                    extracted_text, pages = read_pdf(uploaded_file)
                elif ext == "json":
                    extracted_text = read_json(uploaded_file)
                elif ext in {"html", "htm"}:
                    extracted_text = read_html(uploaded_file)
                else:
                    st.error("Unsupported file type.")
                extracted_text = normalize_text(extracted_text)
                metadata = build_metadata(extracted_text, file_name=file_name, file_type=file_type, rows=rows, pages=pages)
            except Exception as e:
                st.error(f"Could not parse file: {e}")

    with paste_tab:
        pasted = st.text_area("Paste source content", height=260, placeholder="Paste source strings, marketing copy, legal content, support text, HTML, JSON, etc.")
        if pasted and not uploaded_file:
            extracted_text = normalize_text(pasted)
            metadata = build_metadata(extracted_text)

    if extracted_text:
        st.success(f"Content loaded: {metadata.get('word_count', 0)} words, {metadata.get('character_count', 0)} characters.")
        with st.expander("Extracted content preview", expanded=False):
            st.text(extracted_text[:5000] + ("\n...[truncated preview]" if len(extracted_text) > 5000 else ""))

    run = st.button("Run Model-Waldo", type="primary", disabled=not bool(extracted_text))

    if run:
        pair_complexity = language_pair_complexity(source_lang, target_lang)
        sample_words = FAST_SAMPLE_WORDS if fast_mode else MAX_SAMPLE_WORDS
        text_sample = make_sample(extracted_text, max_words=sample_words)
        metadata["llm_sample_word_limit"] = sample_words
        metadata["language_pair_complexity"] = pair_complexity

        with st.spinner("Analyzing source content and scoring providers..."):
            if api_key:
                try:
                    analysis = call_openai_analysis(
                        text_sample=text_sample,
                        metadata=metadata,
                        source_lang=source_lang,
                        target_lang=target_lang,
                        known_domain=selected_content_type,
                        content_requirements=content_requirements,
                        model_name=model_name,
                        api_key=api_key,
                    )
                    analysis_mode = "OpenAI structured analysis"
                except Exception as e:
                    st.error(str(e))
                    st.warning("Falling back to demo rule-based analysis so the app can continue.")
                    analysis = fallback_analysis(extracted_text, metadata, known_domain)
                    analysis_mode = "Fallback rule-based analysis"
            else:
                analysis = fallback_analysis(extracted_text, metadata, known_domain)
                analysis_mode = "Fallback rule-based analysis"

            analysis = apply_content_type_context(analysis, known_domain)
            analysis = apply_content_requirements(analysis, content_requirements)
            analysis = apply_operational_review_policy(analysis, metadata, extracted_text)

            ranked = score_providers(
                providers=providers,
                analysis=analysis,
                source_lang=source_lang,
                target_lang=target_lang,
                priority=priority,
                urgency=urgency,
                pair_complexity=pair_complexity,
                content_requirements=content_requirements,
                metadata=metadata,
            )

        st.caption(f"Analysis mode: {analysis_mode}")

        top = ranked.iloc[0]
        render_decision_dashboard(
            top=top,
            ranked=ranked,
            analysis=analysis,
            pair_complexity=pair_complexity,
            content_requirements=content_requirements,
        )
if dev_inspect_mode:
                render_xai_inspect_panel(
                    analysis=analysis_results, 
                    metadata=metadata, 
                    ranked_df=ranked_providers, 
                    top_provider=top_provider
                )
        render_content_analysis_details(analysis, content_requirements, target_lang)

        with st.expander("5. Model Ranking", expanded=False):
            st.dataframe(ranked[[
                "Provider", "Score", "Quality fit", "Language fit", "Turnaround speed fit", "Terminology fit",
                "Privacy fit", "Workflow type", "Evidence confidence"
            ]], use_container_width=True)
            with st.expander("Evidence notes and limitations", expanded=False):
                st.dataframe(ranked[["Provider", "Evidence notes", "Evidence sources"]], use_container_width=True)
                st.warning(
                    "Provider data is evidence-informed and demo-weighted. Public benchmarks are context-dependent and do not guarantee performance for every domain, language pair, or customer workflow."
                )


if __name__ == "__main__":
    main()
